import os
import re
import base64
import sys
import subprocess
from typing import Union, IO, List
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from sqlalchemy import create_engine
from docx import Document
from supabase import create_client, Client
from supabase.client import ClientOptions
from llama_index.core import SQLDatabase, Settings
from llama_index.llms.ollama import Ollama
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core.query_engine import NLSQLTableQueryEngine
from llama_index.core.prompts.base import PromptTemplate
from llama_index.core.prompts.prompt_type import PromptType


def initialize_supabase() -> Client:
    url = os.environ.get("SUPABASE_URL")
    key = os.environ.get("SUPABASE_KEY")
    
    # print(f"Supabase URL: {url}, Supabase Key: {key}")
    client = create_client(url, key, options=ClientOptions(
        postgrest_client_timeout=100,
        storage_client_timeout=100,
        schema="private"
    ))
    print("Supabase client initialized.")
    return client

def authenticate_gmail():
    SCOPES = ["https://www.googleapis.com/auth/gmail.send"]
    flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
    creds = flow.run_local_server(port=0)
    service = build('gmail', 'v1', credentials=creds)
    print("Gmail authenticated.")
    return service

def initialize_llama() -> NLSQLTableQueryEngine:
    engine = create_engine('postgresql://postgres.SOMETHING@aws-0-us-west-1.pooler.supabase.com:5432/postgres')
    sql_database = SQLDatabase(engine, schema="private", include_tables=["employee"])
    Settings.llm = Ollama(model="llama3", request_timeout=300.0)
    Settings.embed_model = HuggingFaceEmbedding(model_name="BAAI/bge-small-en")
    
    # Initialize the query engine
    query_engine = NLSQLTableQueryEngine(sql_database=sql_database, tables=["employee"], llm=Settings.llm)
    
    # # Define the modified prompt template
    # MO_TEXT_TO_SQL_TMPL = (
    #     "Given the input details, analyze the query to understand what information the user is seeking. "
    #     "If the user has provided specific details such as a name or identifier without asking a direct question, "
    #     "assume they are requesting the email address, and retrieve it accordingly. "
    #     "Create a syntactically correct {dialect} query based on the analysis, execute it, and provide the relevant information. "
    #     "Order the results by a relevant column to return the most useful examples from the database.\n\n"
    #     "Only request a few relevant columns based on the details provided. "
    #     "Be cautious to only query for columns that exist in the schema description. "
    #     "Ensure proper qualification of column names with table names as needed. "
    #     "Follow this format for response:\n\n"
    #     "Input Details: Details here\n"
    #     "SQLQuery: SQL Query to run\n"
    #     "SQLResult: Result of the SQLQuery\n"
    #     "Response: Relevant information here\n\n"
    #     "Use only the tables listed below.\n"
    #     "{schema}\n\n"
    #     "Input Details: {query_str}\n"
    #     "SQLQuery: "
    # )

    # MO_TEXT_TO_SQL_PROMPT = PromptTemplate(
    #     MO_TEXT_TO_SQL_TMPL,
    #     prompt_type=PromptType.TEXT_TO_SQL,
    # )

    # # Update the query engine with the modified prompt template
    # query_engine.update_prompts(
    #     {"sql_retriever:text_to_sql_prompt": MO_TEXT_TO_SQL_PROMPT}
    # )

    print("Llama initialized with modified prompt.")
    return query_engine


def extract_emails(text: str) -> List[str]:
    email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
    emails = email_pattern.findall(text)
    print(f"Extracted emails: {emails}")
    return emails

def get_employee_details(supabase: Client, emails: List[str]) -> List[str]:
    print(f"Fetching employee details for emails: {emails}")
    response = supabase.table('employee').select('designation, department, city').in_('official_email', emails).execute()
    print(f"Supabase response: {response.data}")
    if response.data:
        details_list = [f"{emp['designation']}, {emp['department']}, {emp['city']}" for emp in response.data]
        print(f"Employee details: {details_list}")
        return details_list
    print("No data found for the emails.")
    return ["No data found" for _ in emails]

def replace_placeholders(document: Document, placeholders: List[str], replacements: List[str]) -> None:
    print(f"Replacing placeholders: {placeholders} with {replacements}")
    for para in document.paragraphs:
        for placeholder, replacement in zip(placeholders, replacements):
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, replacement)
    print("Placeholders replaced.")

def populate_table(document: Document, supabase: Client, emails: List[str]) -> None:
    print("Populating table with employee details...")
    table = document.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    employee_details_list = []

    for email in emails:
        print(f"Fetching employee details for email: {email}")
        response = supabase.table('employee').select('designation, department, city').eq('official_email', email).execute()
        print(f"Supabase response for {email}: {response.data}")
        if response.data:
            details = f"{response.data[0]['designation']}, {response.data[0]['department']}, {response.data[0]['city']}"
            employee_details_list.append(details)
        else:
            employee_details_list.append("No data found")
    
    for i, (email, details) in enumerate(zip(emails, employee_details_list), start=1):
        print(f"Adding row for email {email} with details {details}")
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = details
        row_cells[2].text = 'All Items'
    
    print("Table populated.")




def convert_to(folder: str, source: str, output_filename='letter.pdf', timeout=None) -> str:
    print(f"Converting {source} to PDF...")
    args = [libreoffice_exec(), '--headless', '--convert-to', 'pdf', '--outdir', folder, source]
    if not os.path.exists(source):
        raise FileNotFoundError(f"The source file {source} does not exist.")
    if not os.path.exists(folder):
        os.makedirs(folder)
    process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
    if process.returncode != 0:
        raise LibreOfficeError(process.stderr.decode())

    match = re.search(r'-> (.*?) using filter', process.stdout.decode())
    if not match:
        raise LibreOfficeError('Failed to convert the file: No output filename detected.')
    
    generated_filename = match.group(1)
    generated_filepath = os.path.join(folder, generated_filename)
    output_filepath = os.path.join(folder, output_filename)
    
    if os.path.exists(generated_filepath):
        os.rename(generated_filepath, output_filepath)
    else:
        raise LibreOfficeError('Failed to convert the file: Output file does not exist.')
    
    print(f"Conversion successful: {output_filepath}")
    return output_filepath

def libreoffice_exec() -> str:
    if sys.platform == 'darwin':
        return '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    elif sys.platform == 'win32':
        return r'C:\Program Files\LibreOffice\program\soffice.exe'
    elif sys.platform == 'linux':
        return 'libreoffice'
    else:
        raise EnvironmentError('Unsupported platform')

class LibreOfficeError(Exception):
    def __init__(self, output: str):
        super().__init__(output)

def create_message_with_attachment(subject: str, to_emails: str, attachment_path: str) -> dict:
    print(f"Creating message with attachment: {attachment_path}")
    message = MIMEMultipart()
    message['to'] = to_emails
    message['subject'] = subject
    if attachment_path:
        part = MIMEBase('application', 'octet-stream')
        with open(attachment_path, 'rb') as file:
            part.set_payload(file.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(attachment_path)}')
        message.attach(part)
    raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
    print("Message created.")
    return {'raw': raw_message}

def send_message(service, user_id: str, message: dict) -> None:
    print(f"Sending message to {user_id}...")
    try:
        sent_message = service.users().messages().send(userId=user_id, body=message).execute()
        print(f'Sent message to {sent_message["id"]}')
    except HttpError as error:
        print(f'An error occurred: {error}')

def load_document(doc_path: Union[str, IO[bytes], None] = None) -> Document:
    return Document(doc_path)

def process_query(query_str: str) -> dict:
    print(f"Processing query: {query_str}")
    
    query_engine = initialize_llama()
    response = query_engine.query(query_str)
    response_str = str(response)
    llm_emails = extract_emails(response_str)
    return {"status": "success", "emails": llm_emails, "response": response_str}

def create_document(placeholders: List[str], replacements: List[str], emails: List[str]) -> str:
    print(f"Creating document with placeholders: {placeholders} and replacements: {replacements}")
    supabase = initialize_supabase()
    doc_path = "documents/letter_template.docx"
    document = Document(doc_path)
    replace_placeholders(document, placeholders, replacements)
    populate_table(document, supabase, emails)
    modified_doc_path = "files/Letter.docx"
    document.save(modified_doc_path)
    print(f"Document saved as: {modified_doc_path}")
    return modified_doc_path

def send_email_with_attachment(subject: str, to_emails: str, doc_path: str) -> dict:
    print(f"Sending email with attachment: {doc_path}")
    output_folder = 'PDF_directory'
    attachment_path = convert_to(output_folder, doc_path)
    service = authenticate_gmail()
    message = create_message_with_attachment(subject, to_emails, attachment_path)
    send_message(service, "me", message)
    return {"status": "success", "message": "Email sent"}
